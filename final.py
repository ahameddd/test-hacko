import streamlit as st
import PyPDF2
import pandas as pd
import re
import docx
from io import BytesIO
import json
from datetime import datetime
from difflib import SequenceMatcher
import networkx as nx
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Initialize session state for chat history and chat visibility
if "messages" not in st.session_state:
    st.session_state.messages = []
if "current_requirements" not in st.session_state:
    st.session_state.current_requirements = []
if "clarification_questions" not in st.session_state:
    st.session_state.clarification_questions = []
if "version1_requirements" not in st.session_state:
    st.session_state.version1_requirements = []
if "version2_requirements" not in st.session_state:
    st.session_state.version2_requirements = []
if "requirement_estimates" not in st.session_state:
    st.session_state.requirement_estimates = {}
if "requirement_risks" not in st.session_state:
    st.session_state.requirement_risks = {}
if "requirement_dependencies" not in st.session_state:
    st.session_state.requirement_dependencies = {}
if "chat_visible" not in st.session_state:
    st.session_state.chat_visible = False

# Function to extract text from different file types
def extract_text(uploaded_file):
    text = ""
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "pdf":
        try:
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            st.error(f"Error reading PDF file: {e}")
    
    elif file_type in ["txt", "csv"]:
        text = uploaded_file.read().decode("utf-8")
    
    elif file_type == "docx":
        try:
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            st.error(f"Error reading DOCX file: {e}")
    
    else:
        st.error("Unsupported file format. Please upload PDF, DOCX, TXT, or CSV.")
    
    return text

# Function to extract requirements and user stories
def extract_requirements(text):
    requirements = []
    user_stories = []
    capture = False
    
    for line in text.split("\n"):
        line = line.strip()
        
        # Start capturing when we hit Functional/Non-Functional Requirements
        if re.match(r"\d+\.\s*(Functional|Non-Functional) Requirements", line, re.IGNORECASE):
            capture = True
            continue
        
        # Stop capturing when we hit "End of Document" or similar sections
        if re.match(r"\d+\.\s*(Assumptions & Constraints|End of Document)", line, re.IGNORECASE):
            capture = False
            break
        
        if capture and line:
            requirements.append(line)
            if any(kw in line.lower() for kw in ["should", "must", "shall", "require"]):
                user_story = f"As a user, I {line.lower()} so that I can accomplish my goals."
                user_stories.append([user_story])
    
    return "\n".join(requirements), user_stories

# Function to generate clarification questions
def generate_clarification_questions(requirements):
    questions = []
    for req in requirements:
        # Check for ambiguous terms
        if any(term in req.lower() for term in ["should", "may", "can", "could", "might"]):
            questions.append(f"Could you clarify the priority level for: '{req}'?")
        
        # Check for missing details
        if len(req.split()) < 5:
            questions.append(f"Could you provide more details about: '{req}'?")
        
        # Check for technical terms
        if any(term in req.lower() for term in ["system", "interface", "database", "api"]):
            questions.append(f"Could you specify the technical constraints for: '{req}'?")
    
    return questions

# Function to save extracted requirements to a Word document
def save_to_word(requirements_text):
    doc = docx.Document()
    doc.add_heading("Extracted Requirements", level=1)
    doc.add_paragraph(requirements_text)
    
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

# Function to save user stories to an Excel file
def save_to_excel(user_stories):
    df = pd.DataFrame(user_stories, columns=["User Stories"])
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="User Stories", index=False)
    excel_buffer.seek(0)
    return excel_buffer

# Function to compare requirements
def compare_requirements(req1, req2):
    # Calculate similarity ratio between two requirements
    similarity = SequenceMatcher(None, req1.lower(), req2.lower()).ratio()
    
    # Categorize changes based on similarity
    if similarity > 0.9:
        return "Unchanged", similarity
    elif similarity > 0.7:
        return "Minor Changes", similarity
    elif similarity > 0.3:
        return "Major Changes", similarity
    else:
        return "Completely Different", similarity

# Function to estimate cost and effort
def estimate_requirement(req):
    # Simple estimation based on requirement complexity
    words = len(req.split())
    technical_terms = len([term for term in ["system", "interface", "database", "api", "security", "integration"] 
                          if term in req.lower()])
    
    # Calculate effort (in story points)
    effort = min(8, max(1, words // 10 + technical_terms))
    
    # Calculate cost (in thousands)
    cost = effort * 5  # Assuming 5k per story point
    
    # Calculate risk level
    risk_level = "High" if technical_terms > 2 else "Medium" if technical_terms > 0 else "Low"
    
    return {
        "effort": effort,
        "cost": cost,
        "risk_level": risk_level,
        "technical_complexity": technical_terms
    }

# Function to analyze requirement impact
def analyze_impact(req, all_requirements):
    # Create TF-IDF vectorizer
    vectorizer = TfidfVectorizer()
    
    # Create document matrix
    docs = [req] + all_requirements
    tfidf_matrix = vectorizer.fit_transform(docs)
    
    # Calculate cosine similarity
    similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])
    
    # Get top 3 most similar requirements
    top_indices = similarities[0].argsort()[-3:][::-1]
    similar_reqs = [all_requirements[i] for i in top_indices]
    
    return similar_reqs

# Function to create dependency graph
def create_dependency_graph(requirements):
    G = nx.DiGraph()
    
    # Add nodes
    for i, req in enumerate(requirements):
        G.add_node(i, text=req)
    
    # Add edges based on similarity
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(requirements)
    similarities = cosine_similarity(tfidf_matrix)
    
    # Add edges for significant dependencies
    for i in range(len(requirements)):
        for j in range(i + 1, len(requirements)):
            if similarities[i][j] > 0.3:  # Threshold for dependency
                G.add_edge(i, j, weight=similarities[i][j])
    
    return G

# Function to visualize dependency graph
def visualize_dependency_graph(G):
    pos = nx.spring_layout(G)
    
    # Create edges
    edge_trace = go.Scatter(
        x=[], y=[],
        line=dict(width=0.5, color='#888'),
        hoverinfo='none',
        mode='lines')
    
    # Create nodes
    node_trace = go.Scatter(
        x=[], y=[],
        mode='markers+text',
        hoverinfo='text',
        text=[],
        marker=dict(
            showscale=True,
            colorscale='YlOrRd',
            size=10,
            colorbar=dict(
                thickness=15,
                title=dict(
                    text='Node Connections',
                    side='right'
                ),
                xanchor='left'
            )
        ))
    
    # Add edges to trace
    for edge in G.edges():
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        edge_trace['x'] += tuple([x0, x1, None])
        edge_trace['y'] += tuple([y0, y1, None])
    
    # Add nodes to trace
    for node in G.nodes():
        x, y = pos[node]
        node_trace['x'] += tuple([x])
        node_trace['y'] += tuple([y])
        node_trace['text'] += tuple([G.nodes[node]['text']])
    
    # Create the figure
    fig = go.Figure(data=[edge_trace, node_trace],
                   layout=go.Layout(
                       showlegend=False,
                       hovermode='closest',
                       margin=dict(b=20,l=5,r=5,t=40))
                   )
    
    return fig

# Streamlit UI
st.set_page_config(layout="wide")
st.title("Automated Requirement Extraction with AI Assistant")

# Create tabs for different functionalities
tab1, tab2 = st.tabs(["📄 Single Document Analysis", "🔄 Document Comparison"])

# Tab 1 - Original functionality
with tab1:
    # Single column layout for main content
    st.subheader("📄 Document Upload")
    uploaded_file = st.file_uploader("Upload a File", type=["pdf", "docx", "txt", "csv"])
    
    if uploaded_file is not None:
        text = extract_text(uploaded_file)
        
        if text.strip():
            requirements_text, user_stories = extract_requirements(text)
            st.session_state.current_requirements = requirements_text.split("\n")
            
            # Display requirements in an expander
            with st.expander("📋 View Extracted Requirements", expanded=True):
                st.text_area("", requirements_text, height=200)
                
                # Add Analysis section
                st.subheader("📊 Requirement Analysis")
                
                # Create tabs for different analyses
                analysis_tab1, analysis_tab2, analysis_tab3 = st.tabs(["Cost & Effort", "Risk Assessment", "Dependencies"])
                
                with analysis_tab1:
                    st.write("Cost and Effort Estimation")
                    estimates = []
                    for req in st.session_state.current_requirements:
                        estimate = estimate_requirement(req)
                        estimates.append({
                            "Requirement": req,
                            "Effort (Story Points)": estimate["effort"],
                            "Cost (K$)": estimate["cost"],
                            "Technical Complexity": estimate["technical_complexity"]
                        })
                    
                    df_estimates = pd.DataFrame(estimates)
                    
                    # Show summary metrics first
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Estimated Cost", f"${df_estimates['Cost (K$)'].sum():.1f}K")
                    with col2:
                        st.metric("Average Effort", f"{df_estimates['Effort (Story Points)'].mean():.1f} Story Points")
                    with col3:
                        st.metric("Max Complexity", f"{df_estimates['Technical Complexity'].max():.1f}")
                    
                    # Show visualization
                    fig_cost = px.pie(df_estimates, values="Cost (K$)", names="Requirement",
                                    title="Cost Distribution Across Requirements")
                    st.plotly_chart(fig_cost, use_container_width=True)
                    
                    # Show detailed table with pagination
                    st.subheader("Detailed Cost & Effort Table")
                    # Add pagination
                    items_per_page = 5
                    total_pages = len(df_estimates) // items_per_page + (1 if len(df_estimates) % items_per_page > 0 else 0)
                    page = st.number_input("Page", min_value=1, max_value=total_pages, value=1)
                    
                    start_idx = (page - 1) * items_per_page
                    end_idx = min(start_idx + items_per_page, len(df_estimates))
                    paginated_df = df_estimates.iloc[start_idx:end_idx]
                    
                    st.dataframe(paginated_df, use_container_width=True, hide_index=True)
                    st.caption(f"Showing {start_idx + 1}-{end_idx} of {len(df_estimates)} requirements")
                
                with analysis_tab2:
                    st.write("Risk Assessment")
                    risks = []
                    for req in st.session_state.current_requirements:
                        estimate = estimate_requirement(req)
                        risks.append({
                            "Requirement": req,
                            "Risk Level": estimate["risk_level"],
                            "Technical Complexity": estimate["technical_complexity"]
                        })
                    
                    df_risks = pd.DataFrame(risks)
                    
                    # Show risk summary metrics
                    risk_counts = df_risks["Risk Level"].value_counts()
                    col1, col2, col3 = st.columns(3)
                    for i, (risk_level, count) in enumerate(risk_counts.items()):
                        with [col1, col2, col3][i]:
                            st.metric(f"{risk_level} Risk", count)
                    
                    # Show visualization
                    fig_risk = px.pie(df_risks, names="Risk Level", title="Risk Distribution")
                    st.plotly_chart(fig_risk, use_container_width=True)
                    
                    # Show detailed table with pagination
                    st.subheader("Detailed Risk Table")
                    # Add pagination
                    items_per_page = 5
                    total_pages = len(df_risks) // items_per_page + (1 if len(df_risks) % items_per_page > 0 else 0)
                    page = st.number_input("Page", min_value=1, max_value=total_pages, value=1, key="risk_page")
                    
                    start_idx = (page - 1) * items_per_page
                    end_idx = min(start_idx + items_per_page, len(df_risks))
                    paginated_df = df_risks.iloc[start_idx:end_idx]
                    
                    st.dataframe(paginated_df, use_container_width=True, hide_index=True)
                    st.caption(f"Showing {start_idx + 1}-{end_idx} of {len(df_risks)} requirements")
                
                with analysis_tab3:
                    st.write("Requirement Dependencies")
                    G = create_dependency_graph(st.session_state.current_requirements)
                    fig_dep = visualize_dependency_graph(G)
                    st.plotly_chart(fig_dep, use_container_width=True)
                    
                    # Show impact analysis
                    st.subheader("Impact Analysis")
                    selected_req = st.selectbox("Select a requirement to analyze impact:", 
                                              st.session_state.current_requirements)
                    if selected_req:
                        similar_reqs = analyze_impact(selected_req, st.session_state.current_requirements)
                        st.write("Similar/Related Requirements:")
                        for req in similar_reqs:
                            st.markdown(f"- {req}")
            
            # Download buttons in a container
            with st.container():
                st.subheader("📥 Download Options")
                word_file = save_to_word(requirements_text)
                excel_file = save_to_excel(user_stories)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("📄 Download Requirements", word_file, "requirements.docx", 
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                with col2:
                    st.download_button("📊 Download User Stories", excel_file, "user_stories.xlsx", 
                                     "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.warning("No valid text extracted from the uploaded file. Please check the format.")

    # Floating chat button and chat window
    if st.button("💬 Chat with AI Assistant", key="chat_button"):
        st.session_state.chat_visible = not st.session_state.chat_visible
    
    if st.session_state.chat_visible:
        # Create a container for the chat window
        chat_container = st.container()
        with chat_container:
            st.subheader("🤖 AI Assistant")
            
            # Display chat messages
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
            
            # Generate and display clarification questions
            if uploaded_file is not None and text.strip():
                if not st.session_state.clarification_questions:
                    st.session_state.clarification_questions = generate_clarification_questions(st.session_state.current_requirements)
                
                if st.session_state.clarification_questions:
                    with st.expander("💡 Suggested Questions", expanded=True):
                        for question in st.session_state.clarification_questions:
                            if st.button(question):
                                st.session_state.messages.append({"role": "assistant", "content": question})
                                st.rerun()
            
            # Chat input at the bottom
            if prompt := st.chat_input("Ask questions about the requirements"):
                # Add user message to chat history
                st.session_state.messages.append({"role": "user", "content": prompt})
                
                # Generate AI response
                response = f"I understand you're asking about: '{prompt}'. Let me help clarify that requirement."
                st.session_state.messages.append({"role": "assistant", "content": response})
                
                # Rerun to update chat display
                st.rerun()

# Tab 2 - Document Comparison
with tab2:
    st.subheader("🔄 Compare Document Versions")
    
    # Create two columns for document uploads
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("Version 1")
        version1_file = st.file_uploader("Upload Version 1", type=["pdf", "docx", "txt", "csv"], key="v1")
    
    with col2:
        st.write("Version 2")
        version2_file = st.file_uploader("Upload Version 2", type=["pdf", "docx", "txt", "csv"], key="v2")
    
    if version1_file is not None and version2_file is not None:
        # Extract text from both versions
        text1 = extract_text(version1_file)
        text2 = extract_text(version2_file)
        
        if text1.strip() and text2.strip():
            # Extract requirements from both versions
            req1_text, _ = extract_requirements(text1)
            req2_text, _ = extract_requirements(text2)
            
            req1_list = req1_text.split("\n")
            req2_list = req2_text.split("\n")
            
            # Store requirements in session state
            st.session_state.version1_requirements = req1_list
            st.session_state.version2_requirements = req2_list
            
            # Display comparison results
            st.subheader("📊 Comparison Results")
            
            # Create tabs for different comparison views
            comp_tab1, comp_tab2 = st.tabs(["Detailed Comparison", "Summary"])
            
            with comp_tab1:
                # Create a DataFrame for comparison
                comparison_data = []
                for req1 in req1_list:
                    best_match = None
                    best_similarity = 0
                    for req2 in req2_list:
                        status, similarity = compare_requirements(req1, req2)
                        if similarity > best_similarity:
                            best_similarity = similarity
                            best_match = (req2, status, similarity)
                    
                    if best_match:
                        comparison_data.append({
                            "Version 1": req1,
                            "Version 2": best_match[0],
                            "Status": best_match[1],
                            "Similarity": f"{best_match[2]*100:.1f}%"
                        })
                
                # Display comparison table
                df = pd.DataFrame(comparison_data)
                st.dataframe(df, use_container_width=True)
            
            with comp_tab2:
                # Calculate statistics
                total_reqs = len(req1_list)
                unchanged = sum(1 for req in comparison_data if req["Status"] == "Unchanged")
                minor_changes = sum(1 for req in comparison_data if req["Status"] == "Minor Changes")
                major_changes = sum(1 for req in comparison_data if req["Status"] == "Major Changes")
                completely_different = sum(1 for req in comparison_data if req["Status"] == "Completely Different")
                
                # Display summary statistics
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Unchanged", unchanged)
                with col2:
                    st.metric("Minor Changes", minor_changes)
                with col3:
                    st.metric("Major Changes", major_changes)
                with col4:
                    st.metric("Completely Different", completely_different)
                
                # Display pie chart
                fig = px.pie(values=[unchanged, minor_changes, major_changes, completely_different],
                           names=["Unchanged", "Minor Changes", "Major Changes", "Completely Different"],
                           title="Requirement Changes Distribution")
                st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("Please ensure both documents contain valid text.")
